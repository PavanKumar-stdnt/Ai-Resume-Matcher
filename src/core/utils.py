"""
core/utils.py — Multi-format document text extractor + RAG chunking engine.

Supports:
  • PDF  — machine-readable text via pypdf, page-by-page concatenation.
  • DOCX — paragraphs + all table cell text via python-docx.
  • TXT  — raw bytes decoded with a UTF-8 → Latin-1 → UTF-16 fallback chain.

Chunking strategy (sentence-aware sliding window):
  • Splits text into sentences first to avoid cutting mid-sentence.
  • Groups sentences into chunks of ~CHUNK_SIZE words with CHUNK_OVERLAP overlap.
  • Each chunk is tagged with its index and source filename in metadata.
  • Minimum chunk word count enforced to discard noise fragments.

All public functions accept raw bytes so they remain I/O-agnostic and easy to
unit-test without touching the filesystem.
"""

from __future__ import annotations

import io
import logging
import re
from dataclasses import dataclass
from typing import Final

import docx  # python-docx
from pypdf import PdfReader

logger = logging.getLogger(__name__)

# ── Encoding fallback sequence for plain-text files ───────────────────────────
_TXT_ENCODINGS: Final[tuple[str, ...]] = ("utf-8", "latin-1", "utf-16")

# ── Hard cap on raw extracted characters ─────────────────────────────────────
_MAX_CHARS: Final[int] = 50_000

# ── Chunking configuration ────────────────────────────────────────────────────
# Target word count per chunk — ~400 words ≈ 512 tokens for bge-small-en
CHUNK_SIZE: Final[int] = 400
# Overlap in words between consecutive chunks to preserve context continuity
CHUNK_OVERLAP: Final[int] = 80
# Discard chunks shorter than this (noise fragments from sparse pages)
MIN_CHUNK_WORDS: Final[int] = 30


# ─────────────────────────────────────────────────────────────────────────────
# Public data structures
# ─────────────────────────────────────────────────────────────────────────────


@dataclass
class TextChunk:
    """A single chunk produced by the sliding-window chunker."""

    text: str
    chunk_index: int
    word_count: int
    source_filename: str

    def to_metadata(self) -> dict[str, object]:
        """Return a flat dict suitable for Qdrant payload storage."""
        return {
            "chunk_index": self.chunk_index,
            "word_count": self.word_count,
            "source_filename": self.source_filename,
        }


# ─────────────────────────────────────────────────────────────────────────────
# Public API
# ─────────────────────────────────────────────────────────────────────────────


def extract_text_from_bytes(file_bytes: bytes, filename: str) -> str:
    """Dispatch to the correct extractor based on the file extension.

    Parameters
    ----------
    file_bytes : bytes
        Raw binary content of the uploaded file.
    filename : str
        Original filename (used solely to determine format).

    Returns
    -------
    str
        Extracted plain-text content, stripped and capped at _MAX_CHARS.

    Raises
    ------
    ValueError
        If the file extension is not recognised or extraction produces no text.
    """
    lower_name = filename.lower().strip()

    if lower_name.endswith(".pdf"):
        text = _extract_pdf(file_bytes)
    elif lower_name.endswith(".docx"):
        text = _extract_docx(file_bytes)
    elif lower_name.endswith(".txt"):
        text = _extract_txt(file_bytes)
    else:
        ext = lower_name.rsplit(".", 1)[-1] if "." in lower_name else "unknown"
        raise ValueError(
            f"Unsupported file format: '.{ext}'. Accepted: .pdf, .docx, .txt"
        )

    cleaned = _clean_text(text)
    if not cleaned:
        raise ValueError(
            f"No readable text could be extracted from '{filename}'. "
            "The file may be image-only, password-protected, or corrupt."
        )

    logger.debug(
        "Extracted %d chars from '%s' (format=%s).",
        len(cleaned), filename, lower_name.rsplit(".", 1)[-1],
    )
    return cleaned[:_MAX_CHARS]


def chunk_document(text: str, filename: str) -> list[TextChunk]:
    """Split *text* into overlapping sentence-aware sliding-window chunks.

    Algorithm
    ---------
    1. Split the full text into individual sentences using a regex boundary
       detector (handles '.', '!', '?' followed by whitespace/newline).
    2. Accumulate sentences into a word-budget window of CHUNK_SIZE words.
    3. When the budget is reached, emit the chunk, then step back CHUNK_OVERLAP
       words by re-adding trailing sentences to the next window's seed — this
       preserves cross-boundary context (e.g. a skill mentioned at the end of
       one section and the start of the next).
    4. Discard any window with fewer than MIN_CHUNK_WORDS words (sparse noise).

    Parameters
    ----------
    text : str
        Full plain-text content of the document.
    filename : str
        Used to populate metadata on each chunk.

    Returns
    -------
    list[TextChunk]
        Ordered list of non-empty chunks with metadata attached.
    """
    sentences = _split_sentences(text)
    if not sentences:
        return []

    chunks: list[TextChunk] = []
    current_sentences: list[str] = []
    current_word_count: int = 0

    for sentence in sentences:
        sentence_words = len(sentence.split())
        current_sentences.append(sentence)
        current_word_count += sentence_words

        if current_word_count >= CHUNK_SIZE:
            chunk_text = " ".join(current_sentences).strip()
            wc = len(chunk_text.split())
            if wc >= MIN_CHUNK_WORDS:
                chunks.append(TextChunk(
                    text=chunk_text,
                    chunk_index=len(chunks),
                    word_count=wc,
                    source_filename=filename,
                ))

            # Step back by CHUNK_OVERLAP words — keep trailing sentences whose
            # cumulative word count approaches the overlap budget.
            overlap_sentences: list[str] = []
            overlap_wc = 0
            for sent in reversed(current_sentences):
                sw = len(sent.split())
                if overlap_wc + sw > CHUNK_OVERLAP:
                    break
                overlap_sentences.insert(0, sent)
                overlap_wc += sw

            current_sentences = overlap_sentences
            current_word_count = overlap_wc

    # Flush any remaining sentences as the final chunk.
    if current_sentences:
        chunk_text = " ".join(current_sentences).strip()
        wc = len(chunk_text.split())
        if wc >= MIN_CHUNK_WORDS:
            chunks.append(TextChunk(
                text=chunk_text,
                chunk_index=len(chunks),
                word_count=wc,
                source_filename=filename,
            ))

    logger.debug(
        "Chunked '%s' → %d chunks (size=%d, overlap=%d).",
        filename, len(chunks), CHUNK_SIZE, CHUNK_OVERLAP,
    )
    return chunks


# ─────────────────────────────────────────────────────────────────────────────
# Private extractors
# ─────────────────────────────────────────────────────────────────────────────


def _extract_pdf(file_bytes: bytes) -> str:
    """Extract text from a machine-readable PDF using pypdf."""
    parts: list[str] = []
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        if reader.is_encrypted:
            raise ValueError("PDF is password-protected and cannot be read.")
        for page_index, page in enumerate(reader.pages):
            try:
                parts.append(page.extract_text() or "")
            except Exception as exc:  # noqa: BLE001
                logger.warning("PDF page %d extraction failed: %s", page_index, exc)
    except ValueError:
        raise
    except Exception as exc:
        raise ValueError(f"Failed to parse PDF: {exc}") from exc
    return "\n".join(parts)


def _extract_docx(file_bytes: bytes) -> str:
    """Extract text from a Word document (.docx) via python-docx."""
    parts: list[str] = []
    try:
        document = docx.Document(io.BytesIO(file_bytes))
        for paragraph in document.paragraphs:
            t = paragraph.text.strip()
            if t:
                parts.append(t)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    t = cell.text.strip()
                    if t:
                        parts.append(t)
    except Exception as exc:
        raise ValueError(f"Failed to parse DOCX: {exc}") from exc
    return "\n".join(parts)


def _extract_txt(file_bytes: bytes) -> str:
    """Decode a plain-text file using a UTF-8 → Latin-1 → UTF-16 fallback chain."""
    for encoding in _TXT_ENCODINGS:
        try:
            text = file_bytes.decode(encoding)
            logger.debug("TXT decoded with encoding '%s'.", encoding)
            return text
        except (UnicodeDecodeError, LookupError):
            logger.debug("TXT decoding failed with '%s', trying next.", encoding)
    raise ValueError(
        f"Could not decode .txt with any of: {', '.join(_TXT_ENCODINGS)}."
    )


# ─────────────────────────────────────────────────────────────────────────────
# Text processing helpers
# ─────────────────────────────────────────────────────────────────────────────


def _clean_text(text: str) -> str:
    """Normalise whitespace while preserving meaningful line breaks."""
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[^\S\n]+", " ", text)
    return text.strip()


def _split_sentences(text: str) -> list[str]:
    """Split text into sentences using a punctuation-aware regex.

    Handles common resume patterns: bullet points, numbered lists, and
    sentences ending with '.', '!', '?' followed by whitespace or newlines.
    Newlines that don't follow sentence-ending punctuation are treated as
    soft separators (section headers, bullet items).
    """
    # Treat newlines as sentence boundaries in addition to punctuation.
    normalised = re.sub(r"\n+", " . ", text)
    # Split on sentence-ending punctuation followed by whitespace + capital or digit.
    raw = re.split(r"(?<=[.!?])\s+(?=[A-Z0-9\"\'])", normalised)
    sentences: list[str] = []
    for s in raw:
        s = s.strip()
        if s:
            sentences.append(s)
    return sentences
